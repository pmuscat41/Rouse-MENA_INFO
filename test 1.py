# Create a polished .docx with one-pager sections and formatted tables
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from docx.oxml import OxmlElement

# --------- helpers ---------
def set_cell_margins(cell, **kwargs):
    """
    cell:  actual cell instance you want to modify
    interface:
        set_cell_margins(cell, top=50, start=50, bottom=50, end=50)
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m in ['top', 'start', 'bottom', 'end']:
        if m in kwargs:
            node = OxmlElement(f'w:{m}')
            node.set(qn('w:w'), str(kwargs.get(m)))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
    tcPr.append(tcMar)

def shade_row(row, color="D9D9D9"):
    """Shade a table row with a hex color (no #)."""
    for cell in row.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), color)
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        tcPr.append(shd)

def set_col_widths(table, widths_in_inches):
    for row in table.rows:
        for cell, width in zip(row.cells, widths_in_inches):
            cell.width = Inches(width)

def add_keyfacts_table(container_cell, facts):
    t = container_cell.add_table(rows=0, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.style = 'Table Grid'
    for label, value in facts:
        r = t.add_row()
        r.cells[0].text = str(label)
        r.cells[0].paragraphs[0].runs[0].font.bold = True
        r.cells[1].text = str(value)
    # style tweaks
    for row in t.rows:
        for c in row.cells:
            set_cell_margins(c, top=60, bottom=60, start=80, end=80)
    try:
        set_col_widths(t, [2.0, 4.3])
    except:
        pass
    return t

def add_requirements_table(doc, rows):
    # 6 columns: Document, Required? (P/D), Deadline from Filing, Form (what is filed), Execution (how), Notes
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Light List Accent 1' if 'Light List Accent 1' in [s.name for s in doc.styles] else 'Table Grid'
    hdr = table.rows[0].cells
    headers = ["Document", "Required? (P/D)", "Deadline from Filing", "Form (filed)", "Execution (how)", "Notes"]
    for i, h in enumerate(headers):
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_margins(hdr[i], top=80, bottom=120, start=80, end=80)
    shade_row(table.rows[0], color="EDEDED")

    for r in rows:
        row = table.add_row()
        cells = row.cells
        cells[0].text = r.get("Document","")
        cells[1].text = r.get("Required","")
        cells[2].text = r.get("Deadline","")
        cells[3].text = r.get("Form","")
        cells[4].text = r.get("Execution","")
        cells[5].text = r.get("Notes","")
        for c in cells:
            for run in c.paragraphs[0].runs:
                run.font.size = Pt(10.5)
            set_cell_margins(c, top=60, bottom=60, start=80, end=80)

    # Column widths (approx, in inches) to fit typical page with modest margins
    try:
        set_col_widths(table, [1.6, 0.9, 1.1, 1.2, 1.2, 1.7])
    except:
        pass
    return table

def add_country_section(doc, country_name, map_path_placeholder, population, system, office, languages, translation_rule, req_rows, notes=None):
    # Country Heading
    h = doc.add_heading(country_name + " — Patents & Designs", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Top layout: 1 row x 2 cols table (left: map placeholder, right: key facts)
    wrap = doc.add_table(rows=1, cols=2)
    wrap.alignment = WD_TABLE_ALIGNMENT.CENTER
    wrap.style = 'Table Grid'
    left = wrap.rows[0].cells[0]
    right = wrap.rows[0].cells[1]

    # Left: map placeholder box
    p = left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"[Map placeholder: {map_path_placeholder}]")
    r.italic = True
    r.font.size = Pt(10.5)

    set_cell_margins(left, top=200, bottom=200, start=80, end=80)
    set_cell_margins(right, top=120, bottom=120, start=80, end=80)

    # Right: key facts table
    facts = [
        ("Population", population),
        ("System / Coverage", system),
        ("Office / Authority", office),
        ("Filing language(s)", languages),
        ("Translation rule", translation_rule),
        ("Last verified", "2025‑08‑09"),
    ]
    add_keyfacts_table(right, facts)

    # Documentary Requirements
    doc.add_paragraph()  # spacing
    sub = doc.add_paragraph("Documentary Requirements (Patents & Designs)")
    sub.runs[0].bold = True

    add_requirements_table(doc, req_rows)

    if notes:
        doc.add_paragraph()
        n = doc.add_paragraph("Notes: ")
        n.runs[0].bold = True
        doc.add_paragraph(notes)

# --------- document build ---------
document = Document()

# Adjust base styles
normal_style = document.styles['Normal']
normal_style.font.name = 'Calibri'
normal_style.font.size = Pt(11)

# Set narrower margins for better table fit
for section in document.sections:
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

# Cover
title = document.add_paragraph()
title_run = title.add_run("Rouse MEA — Patents & Designs: Country One‑Pagers")
title_run.bold = True
title_run.font.size = Pt(18)
title.alignment = WD_ALIGN_PARAGRAPH.LEFT

subtitle = document.add_paragraph("Version 1.0  ·  Prepared: 2025‑08‑09")
subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
document.add_paragraph("Scope: Minimal one‑page summaries per jurisdiction for patent & design filings. Tables capture document requirements, deadlines, and execution levels.")

document.add_paragraph("")

# Data for each jurisdiction
countries = []

# UAE
countries.append({
    "name": "United Arab Emirates (UAE)",
    "map": "maps/uae.png",
    "population": "10,876,981 (2024)",
    "system": "National (Paris/PCT NP)",
    "office": "UAE Ministry of Economy — Patent Office",
    "languages": "Arabic & English",
    "translation": "Arabic translation as required (title/abstract at filing; full text by invitation/≈90 days)",
    "reqs": [
        {"Document":"Power of Attorney","Required":"Y/Y","Deadline":"90 days","Form":"Original","Execution":"Notarized","Notes":"Practice varies re: legalization; confirm current rule. Conf ~70%."},
        {"Document":"Assignment (inventor→applicant)","Required":"Y/Y","Deadline":"90 days","Form":"Original","Execution":"Notarized","Notes":"If applicant ≠ inventor. Conf ~70%."},
        {"Document":"Commercial/Trade Licence or Register Extract","Required":"Y/Y","Deadline":"90 days","Form":"Copy or certified copy","Execution":"Often notarized/legalized","Notes":"Corporate applicants. Conf ~60%."},
        {"Document":"Priority Document (if claimed)","Required":"Y/Y","Deadline":"3 months","Form":"Certified copy","Execution":"Simple","Notes":"PCT NP usually via IB."},
    ],
    "notes": "Missing mandatory docs typically rectifiable within 90 days; late fees/office invitation may apply."
})

# Saudi Arabia
countries.append({
    "name": "Saudi Arabia (KSA)",
    "map": "maps/saudi.png",
    "population": "≈33.9 million (2024 est.)",
    "system": "National (Paris/PCT NP)",
    "office": "Saudi Authority for Intellectual Property (SAIP)",
    "languages": "Arabic & English",
    "translation": "Arabic components as required; English specs accepted with translations if invited",
    "reqs": [
        {"Document":"Power of Attorney","Required":"Y/Y","Deadline":"Up to 3 months (by invitation)","Form":"Original","Execution":"Apostille or Consular legalization","Notes":"Apostille accepted. Conf ~90%."},
        {"Document":"Assignment (inventor→applicant)","Required":"Y/Y","Deadline":"Up to 3 months","Form":"Original","Execution":"Apostille or Consular legalization","Notes":"If applicant ≠ inventor. Conf ~85%."},
        {"Document":"Commercial/Certificate of Incorporation","Required":"N/N (routine)","Deadline":"—","Form":"—","Execution":"—","Notes":"Not routinely required. Conf ~70%."},
        {"Document":"Priority Document (if claimed)","Required":"Y/Y","Deadline":"~60–90 days","Form":"Certified copy","Execution":"Simple","Notes":"Translation if requested."},
    ],
    "notes": "Deadlines generally triggered by SAIP invitation; docket buffer recommended."
})

# Oman
countries.append({
    "name": "Oman",
    "map": "maps/oman.png",
    "population": "5,049,269 (2023)",
    "system": "National (Paris/PCT NP)",
    "office": "Ministry of Commerce, Industry & Investment Promotion (IP Department)",
    "languages": "Arabic & English",
    "translation": "Arabic translation required per office timelines",
    "reqs": [
        {"Document":"Power of Attorney","Required":"Y/Y","Deadline":"60 days","Form":"Original","Execution":"Apostille or Consular legalization","Notes":"Hard deadline; app may lapse if missed. Conf ~90%."},
        {"Document":"Assignment (inventor→applicant)","Required":"Y/Y","Deadline":"≤90 days","Form":"Original","Execution":"Apostille or Consular legalization","Notes":"Conf ~85%."},
        {"Document":"Commercial/Certificate of Incorporation","Required":"Y/Y","Deadline":"≤90 days","Form":"Certified copy","Execution":"Apostille or Consular legalization","Notes":"Arabic translation typically required. Conf ~85%."},
        {"Document":"Priority Document (if claimed)","Required":"Y/Y","Deadline":"3 months","Form":"Certified copy","Execution":"Simple","Notes":"No translation generally required. Conf ~80%."},
    ],
    "notes": "Ensure Arabic translations are scheduled early."
})

# Qatar
countries.append({
    "name": "Qatar",
    "map": "maps/qatar.png",
    "population": "2,979,082 (2023)",
    "system": "National (Paris/PCT NP)",
    "office": "Ministry of Commerce & Industry — Industrial Property Office",
    "languages": "Arabic & English",
    "translation": "Key parts in Arabic; translation window up to ~6 months in practice",
    "reqs": [
        {"Document":"Power of Attorney","Required":"Y/Y","Deadline":"Up to 6 months","Form":"Original","Execution":"Consular legalization (Qatar not in Apostille)","Notes":"Conf ≥90%."},
        {"Document":"Assignment (if applicant changed post‑PCT)","Required":"If applicable","Deadline":"Within 6 months/by invitation","Form":"Original","Execution":"Consular legalization","Notes":"Conf ~80%."},
        {"Document":"Commercial/Certificate of Incorporation","Required":"N/N (routine)","Deadline":"—","Form":"—","Execution":"—","Notes":"Not typically required."},
        {"Document":"Priority Document (if claimed)","Required":"Y/Y","Deadline":"3 months (practice up to 6)","Form":"Certified copy","Execution":"Simple","Notes":"Translation if requested."},
    ],
    "notes": "File Arabic title/abstract/claims at filing; full spec timing per office practice."
})

# Bahrain
countries.append({
    "name": "Bahrain",
    "map": "maps/bahrain.png",
    "population": "1,588,670 (2024)",
    "system": "National (Paris/PCT NP)",
    "office": "Ministry of Industry & Commerce — Industrial Property Office",
    "languages": "Arabic & English",
    "translation": "Arabic translation per office timelines",
    "reqs": [
        {"Document":"Power of Attorney","Required":"Y/Y","Deadline":"3 months","Form":"Original","Execution":"Apostille or Consular legalization","Notes":"Conf ≥85%."},
        {"Document":"Assignment (inventor→applicant)","Required":"Y/Y","Deadline":"3 months","Form":"Original","Execution":"Apostille or Consular legalization","Notes":"Conf ≥85%."},
        {"Document":"Commercial/Corporate proof","Required":"Y/Y","Deadline":"3 months","Form":"Certified copy","Execution":"Legalized","Notes":"Corporate applicants."},
        {"Document":"Priority Document (if claimed)","Required":"Y/Y","Deadline":"3 months","Form":"Certified copy","Execution":"Simple","Notes":"Translation if required."},
    ],
    "notes": "Allow time for legalization and couriering originals."
})

# Kuwait
countries.append({
    "name": "Kuwait",
    "map": "maps/kuwait.png",
    "population": "4,973,861 (2024)",
    "system": "National (Paris/PCT NP)",
    "office": "Kuwait Patent Office (KPO)",
    "languages": "Arabic & English",
    "translation": "Arabic translation per office practice",
    "reqs": [
        {"Document":"Power of Attorney","Required":"Y/Y","Deadline":"90 days","Form":"Original","Execution":"Consular legalization","Notes":"Copy at filing; original follows. Conf ≥85%."},
        {"Document":"Assignment (inventor→applicant)","Required":"Y/Y","Deadline":"90 days","Form":"Original","Execution":"Consular legalization","Notes":"Conf ≥85%."},
        {"Document":"Commercial/Register Extract","Required":"Y/Y","Deadline":"90 days","Form":"Certified copy","Execution":"Consular legalization","Notes":"Corporate applicants."},
        {"Document":"Priority Document (if claimed)","Required":"Y/Y","Deadline":"~90 days","Form":"Certified copy","Execution":"Simple","Notes":"Translation if requested."},
    ],
    "notes": "Docket early for consular timeframes."
})

# Egypt
countries.append({
    "name": "Egypt",
    "map": "maps/egypt.png",
    "population": "116,538,258 (2024)",
    "system": "National (Paris/PCT NP)",
    "office": "Egyptian Patent Office (EGPO), ASRT",
    "languages": "Arabic (English accepted with Arabic translation)",
    "translation": "Arabic translation is mandatory within office timelines",
    "reqs": [
        {"Document":"Power of Attorney","Required":"Y/Y","Deadline":"≈4 months","Form":"Original","Execution":"Consular legalization (Egyptian Consulate)","Notes":"Conf ≥85%."},
        {"Document":"Assignment (inventor→applicant)","Required":"Y/Y","Deadline":"≈4 months","Form":"Original","Execution":"Consular legalization","Notes":"Conf ≥85%."},
        {"Document":"Commercial/Certificate of Incorporation","Required":"Y/Y","Deadline":"≈4 months","Form":"Certified copy","Execution":"Consular legalization","Notes":"Arabic translation required."},
        {"Document":"Priority Document (if claimed)","Required":"Y/Y","Deadline":"~3 months","Form":"Certified copy","Execution":"Simple + Arabic translation if needed","Notes":"Conf ≥85%."},
    ],
    "notes": "Strict formalities windows; missing documents risk abandonment."
})

# South Africa
countries.append({
    "name": "South Africa (RSA)",
    "map": "maps/south_africa.png",
    "population": "≈63.1 million (2024 est.)",
    "system": "National (Paris/PCT NP)",
    "office": "Companies and Intellectual Property Commission (CIPC)",
    "languages": "English",
    "translation": "No translation required (English filing)",
    "reqs": [
        {"Document":"Power of Attorney (Form P3)","Required":"Y/Y","Deadline":"Within ~6 months (extendable)","Form":"Original or signed copy","Execution":"Simple (no notarization/legalization)","Notes":"Conf ≥85%."},
        {"Document":"Proof of entitlement / Assignment","Required":"Y/Y (if applicant ≠ inventor)","Deadline":"Within ~6 months (pre‑acceptance)","Form":"Signed assignment/affidavit","Execution":"Simple","Notes":"Conf ≥85%."},
        {"Document":"Commercial/Corporate docs","Required":"N/N","Deadline":"—","Form":"—","Execution":"—","Notes":"Not required."},
        {"Document":"Priority Document (if Paris)","Required":"Y/Y","Deadline":"Within 6 months (extendable)","Form":"Certified copy","Execution":"Simple","Notes":"PCT NP handled via IB."},
    ],
    "notes": "Deposit system (formalities); docket acceptance deadlines."
})

# Nigeria
countries.append({
    "name": "Nigeria",
    "map": "maps/nigeria.png",
    "population": "232,679,478 (2024)",
    "system": "National (Paris/PCT NP)",
    "office": "Trademarks, Patents & Designs Registry (FMITI)",
    "languages": "English",
    "translation": "No translation required (English filing)",
    "reqs": [
        {"Document":"Power of Attorney","Required":"Y/Y","Deadline":"At/soon after filing","Form":"Signed copy/original","Execution":"Simple (no notarization)","Notes":"Conf ~75%."},
        {"Document":"Assignment (inventor→applicant)","Required":"Y/Y (good practice)","Deadline":"Prompt / by invitation","Form":"Signed assignment","Execution":"Simple","Notes":"Not always required to proceed."},
        {"Document":"Commercial/Corporate docs","Required":"N/N (routine)","Deadline":"—","Form":"—","Execution":"—","Notes":"Not typically required."},
        {"Document":"Priority Document (if Paris)","Required":"Y/Y","Deadline":"With/shortly after filing","Form":"Certified copy","Execution":"Simple","Notes":"PCT NP via IB."},
    ],
    "notes": "Non‑examining system; originals may be invited before grant."
})

# Kenya
countries.append({
    "name": "Kenya",
    "map": "maps/kenya.png",
    "population": "56,432,944 (2024)",
    "system": "National (Paris/PCT NP) or ARIPO designating KE",
    "office": "Kenya Industrial Property Institute (KIPI)",
    "languages": "English",
    "translation": "No translation required (English filing)",
    "reqs": [
        {"Document":"Power of Attorney","Required":"Y/Y","Deadline":"At/after filing (on invitation)","Form":"Signed copy/original","Execution":"Simple","Notes":"Conf ~80%."},
        {"Document":"Assignment (inventor→applicant)","Required":"Y/Y (if applicant ≠ inventor)","Deadline":"Before grant/recordal","Form":"Signed assignment","Execution":"Simple","Notes":"Conf ~80%."},
        {"Document":"Commercial/Corporate docs","Required":"N/N (routine)","Deadline":"—","Form":"—","Execution":"—","Notes":"Not required."},
        {"Document":"Priority Document (if Paris)","Required":"Y/Y","Deadline":"Upon request","Form":"Certified copy","Execution":"Simple","Notes":"PCT NP via IB."},
    ],
    "notes": "Substantive examination applies; keep PoA on file early."
})

# ARIPO
countries.append({
    "name": "ARIPO (AP) — Regional",
    "map": "maps/aripo.png",
    "population": "— (regional system; 21 member states)",
    "system": "Regional (Paris/PCT regional phase)",
    "office": "African Regional Intellectual Property Organization (Harare)",
    "languages": "English (working language)",
    "translation": "English required; translate priority docs if non‑EN",
    "reqs": [
        {"Document":"Power of Attorney (Form 4)","Required":"Y/Y","Deadline":"If not at entry, within ~2 months after entry","Form":"Signed PoA","Execution":"Simple (no legalization)","Notes":"Conf ≥85%."},
        {"Document":"Assignment / proof of entitlement","Required":"If applicant ≠ inventor","Deadline":"By invitation / ~4 months","Form":"Signed assignment","Execution":"Simple","Notes":"Conf ≥80%."},
        {"Document":"Commercial/Corporate docs","Required":"N/N","Deadline":"—","Form":"—","Execution":"—","Notes":"Not required."},
        {"Document":"Priority Document (if Paris)","Required":"Y/Y","Deadline":"3 months","Form":"Certified copy","Execution":"Simple (+ EN translation if needed)","Notes":"Conf ≥85%."},
    ],
    "notes": "Grant effective in designated member states without further validation."
})

# OAPI
countries.append({
    "name": "OAPI (OA) — Regional",
    "map": "maps/oapi.png",
    "population": "— (regional system; 17 member states)",
    "system": "Regional (Paris; direct regional filing)",
    "office": "Organisation Africaine de la Propriété Intellectuelle (Yaoundé)",
    "languages": "French (official)",
    "translation": "French required; translate priority docs if non‑FR/EN",
    "reqs": [
        {"Document":"Power of Attorney","Required":"Y/Y","Deadline":"At filing or short window (~3 months)","Form":"Signed PoA","Execution":"Simple (no authentication)","Notes":"Conf ~75%."},
        {"Document":"Assignment / proof of entitlement","Required":"If applicant ≠ inventor","Deadline":"Soon after filing","Form":"Original or certified copy","Execution":"Simple","Notes":"Conf ~75%."},
        {"Document":"Commercial/Corporate docs","Required":"N/N","Deadline":"—","Form":"—","Execution":"—","Notes":"Not required."},
        {"Document":"Priority Document (if Paris)","Required":"Y/Y","Deadline":"3–6 months","Form":"Certified copy","Execution":"Simple (+ FR/EN translation if needed)","Notes":"Conf ~75%."},
    ],
    "notes": "Single filing covers all member states; observe strict priority‑doc timeline."
})

# Build sections
for idx, c in enumerate(countries):
    add_country_section(
        document, c["name"], c["map"], c["population"], c["system"], c["office"],
        c["languages"], c["translation"], c["reqs"], c.get("notes")
    )
    if idx != len(countries) - 1:
        document.add_page_break()

# Footer note
for section in document.sections:
    footer = section.footer
    p = footer.paragraphs[0]
    p.text = "Prepared 2025‑08‑09 · Verify formalities with latest office circulars/local counsel before filing."
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Save file
out_path = "/mnt/data/Rouse_MEA_Patents_Designs_OnePagers_v1.docx"
document.save(out_path)
out_path
